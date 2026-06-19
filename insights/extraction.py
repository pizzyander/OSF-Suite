"""
extraction.py — extract plain text from PDF, DOCX, or raw text.

Dependencies (add to your insights/requirements.txt):
    pypdf>=3.0.0
    python-docx>=1.0.0
"""

import io
from typing import Tuple


SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".text"}
MAX_TEXT_LENGTH = 50_000  # characters — keeps prompt injection reasonable


class ExtractionError(ValueError):
    """Raised when a file cannot be parsed or yields no usable text."""
    pass


def extract_from_pdf(data: bytes) -> str:
    try:
        from pypdf import PdfReader
    except ImportError:
        raise ExtractionError("pypdf is not installed. Add pypdf>=3.0.0 to requirements.txt")

    try:
        reader = PdfReader(io.BytesIO(data))
    except Exception as e:
        raise ExtractionError(f"Could not open PDF: {e}")

    pages = []
    for i, page in enumerate(reader.pages):
        try:
            text = page.extract_text() or ""
            pages.append(text.strip())
        except Exception:
            continue

    full_text = "\n\n".join(p for p in pages if p)
    if not full_text.strip():
        raise ExtractionError(
            "PDF appears to be scanned/image-only — no extractable text found. "
            "Please upload a text-based PDF or paste the content manually."
        )
    return full_text


def extract_from_docx(data: bytes) -> str:
    try:
        from docx import Document
    except ImportError:
        raise ExtractionError("python-docx is not installed. Add python-docx>=1.0.0 to requirements.txt")

    try:
        doc = Document(io.BytesIO(data))
    except Exception as e:
        raise ExtractionError(f"Could not open DOCX: {e}")

    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

    # Also extract text from tables (pricing cards often live in tables)
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                paragraphs.append(row_text)

    full_text = "\n\n".join(paragraphs)
    if not full_text.strip():
        raise ExtractionError("DOCX file appears to be empty or contains no readable text.")
    return full_text


def extract_from_text(data: bytes) -> str:
    for encoding in ("utf-8", "utf-16", "latin-1"):
        try:
            return data.decode(encoding).strip()
        except (UnicodeDecodeError, ValueError):
            continue
    raise ExtractionError("Text file could not be decoded — please use UTF-8 encoding.")


def extract(data: bytes, filename: str) -> Tuple[str, str]:
    """
    Extract plain text from uploaded bytes.

    Args:
        data:     raw file bytes
        filename: original filename, used to detect format

    Returns:
        (extracted_text, source_type)  where source_type is "pdf"|"docx"|"text"

    Raises:
        ExtractionError: if the format is unsupported or parsing fails
    """
    name_lower = filename.lower()

    if name_lower.endswith(".pdf"):
        source_type = "pdf"
        text = extract_from_pdf(data)
    elif name_lower.endswith(".docx"):
        source_type = "docx"
        text = extract_from_docx(data)
    elif name_lower.endswith((".txt", ".text")):
        source_type = "text"
        text = extract_from_text(data)
    else:
        ext = "." + name_lower.rsplit(".", 1)[-1] if "." in name_lower else "(none)"
        raise ExtractionError(
            f"Unsupported file type '{ext}'. Please upload a PDF, DOCX, or plain text file."
        )

    if len(text) > MAX_TEXT_LENGTH:
        text = text[:MAX_TEXT_LENGTH] + "\n\n[Document truncated to 50,000 characters]"

    return text.strip(), source_type


def extract_from_raw_text(raw: str) -> Tuple[str, str]:
    """
    Accept a plain string (from a JSON/form text field) as context.
    Returns (text, "text").
    """
    text = raw.strip()
    if not text:
        raise ExtractionError("Text input is empty.")
    if len(text) > MAX_TEXT_LENGTH:
        text = text[:MAX_TEXT_LENGTH] + "\n\n[Input truncated to 50,000 characters]"
    return text, "text"